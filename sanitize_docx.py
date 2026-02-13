import os
import json
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from openai import AzureOpenAI

load_dotenv()

# Azure OpenAI config
client = AzureOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2024-02-01"  # Update to latest stable version if needed (check docs)
)

deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")


def remove_images_and_add_placeholders(doc: Document):
    """
    Removes inline images from the main document body
    and replaces each one with a numbered placeholder like [Photo-1].
    Also handles headers and footers (common for logos).
    """
    photo_counter = 1

    # 1. Main body paragraphs
    for paragraph in doc.paragraphs:
        i = 0
        while i < len(paragraph.runs):
            run = paragraph.runs[i]
            
            drawing = run._element.find(qn('w:drawing'))
            if drawing is not None:
                run.clear()
                
                placeholder = f"[Photo-{photo_counter}]"
                photo_counter += 1
                
                new_run = paragraph.add_run(placeholder + " ")
                new_run.bold = True
                
                i += 1
            else:
                i += 1

    # 2. Headers & Footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            i = 0
            while i < len(header.runs):
                run = header.runs[i]
                if run._element.find(qn('w:drawing')) is not None:
                    run.clear()
                    header.add_run(f"[Header Image] ").bold = True
                i += 1

        for footer in section.footer.paragraphs:
            i = 0
            while i < len(footer.runs):
                run = footer.runs[i]
                if run._element.find(qn('w:drawing')) is not None:
                    run.clear()
                    footer.add_run(f"[Footer Image] ").bold = True
                i += 1

    print(f"→ Replaced {photo_counter - 1} image(s) with placeholders")
    return doc


def sanitize_text_block(text: str, para_or_cell) -> bool:
    """
    Shared logic: send text to Azure OpenAI → get replacements → apply them
    Returns True if any replacement was made
    """
    if not text.strip():
        return False

    prompt = f"""
Analyze the following text for sensitive information such as:
- full names, first/last names
- emails
- phone numbers (any format)
- addresses (street, city, postal code)
- financial details (account numbers, credit cards, amounts)
- identity numbers (NIC, passport, employee ID)
- dates of birth

For each identified piece of sensitive data:
- provide the EXACT original string to replace
- suggest a context-appropriate dummy replacement
- classify the type

Output ONLY valid JSON — a list of objects.
If no PII found, return empty list [].

Example output format:
[{{"original": "john.doe87@company.lk", "type": "email", "replacement": "user@example.com"}},
 {{"original": "+94 77 123 4567", "type": "phone", "replacement": "+94 77 000 0000"}}]

Text:
{text}
"""

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "You are a precise PII detection and anonymization assistant. Return only clean JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=1500,
        )

        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content.split("```json")[1].split("```")[0].strip()

        replacements = json.loads(content)

        if replacements:
            new_text = text
            for rep in replacements:
                new_text = new_text.replace(rep["original"], rep["replacement"])
            para_or_cell.text = new_text
            return True

        return False

    except json.JSONDecodeError:
        print("    → JSON parse error — skipping")
        return False
    except Exception as e:
        print(f"    → Error: {e}")
        return False


def sanitize_docx(input_path, output_path):
    print(f"Processing: {input_path}")

    # Load .docx
    doc = Document(input_path)

    # Step 1: Remove images and insert placeholders
    doc = remove_images_and_add_placeholders(doc)

    # Step 2: Sanitize normal paragraphs
    paragraph_count = 0
    replaced_paras = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        paragraph_count += 1
        print(f"  Paragraph {paragraph_count} ({len(text)} chars)")

        if sanitize_text_block(para.text, para):
            replaced_paras += 1

    # Step 3: Sanitize tables
    table_count = 0
    replaced_cells = 0
    for table_idx, table in enumerate(doc.tables, 1):
        table_count += 1
        print(f"  Table {table_idx} ({len(table.rows)} rows)")

        for row_idx, row in enumerate(table.rows, 1):
            for cell_idx, cell in enumerate(row.cells, 1):
                # A cell can contain multiple paragraphs
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    print(f"    Cell R{row_idx}C{cell_idx} ({len(text)} chars): {text[:60]}...")

                    if sanitize_text_block(para.text, para):
                        replaced_cells += 1

    print(f"\nSummary:")
    print(f"  • Processed {paragraph_count} paragraphs → replaced in {replaced_paras}")
    print(f"  • Processed {table_count} tables → replaced in {replaced_cells} cells")
    print(f"\nSanitized document saved to:\n  {output_path}")

    doc.save(output_path)


# ───────────────────────────────────────────────
#  Local execution
# ───────────────────────────────────────────────
if __name__ == "__main__":
    input_file = "input.docx"           # ← your original file
    output_file = "sanitized_complex_report.docx"

    if not os.path.exists(input_file):
        print(f"Error: Input file not found → {input_file}")
    else:
        sanitize_docx(input_file, output_file)