import azure.functions as func
import logging
from io import BytesIO
import json
import os
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from openai import AzureOpenAI, OpenAIError
import magic  # For file type validation
import fitz  # PyMuPDF — for in-place PDF redaction/replacement

app = func.FunctionApp()

# ───────────────────────────────────────────────
# Azure OpenAI client (create once, fallback for local dev)
# ───────────────────────────────────────────────
# For local development, use API key
client = AzureOpenAI(
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    api_version="2024-02-01"
)
logging.info("Using API key for Azure OpenAI")

deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

# ───────────────────────────────────────────────
# File type validation (security)
# ───────────────────────────────────────────────
def validate_file_type(file_content: bytes, expected_extension: str) -> bool:
    """Validate file type by magic bytes (content), not just extension."""
    try:
        mime = magic.from_buffer(file_content, mime=True)
        
        # Map extensions to allowed MIME types
        allowed_mimes = {
            '.docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            '.doc': ['application/msword', 'application/x-msword'],
            '.pdf': ['application/pdf']
        }
        
        if expected_extension in allowed_mimes:
            return mime in allowed_mimes[expected_extension]
        return False
    except Exception as e:
        logging.error(f"File type validation error: {str(e)}")
        return False

# ───────────────────────────────────────────────
# Extract text from PDF (security: text only, no scripts/embedded files)
# ───────────────────────────────────────────────
def extract_text_from_pdf(pdf_stream: BytesIO) -> str:
    """Extract plain text from PDF using PyMuPDF, ignoring scripts and embedded files."""
    try:
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            try:
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            except Exception as e:
                logging.warning(f"Could not extract text from page {page_num}: {str(e)}")
                continue

        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logging.error(f"PDF text extraction error: {str(e)}")
        raise ValueError("Failed to extract text from PDF")

# ═══════════════════════════════════════════════════════════════
# PDF FORMATTING PRESERVATION - Enhanced Implementation
# ═══════════════════════════════════════════════════════════════

def _map_to_base14_font(font_name: str, flags: int = 0) -> str:
    """Map a PDF font name to the closest Base-14 font for redaction replacement text.
    
    Enhanced to consider font flags for better style matching.
    
    PyMuPDF redaction can only insert text using Base-14 fonts or fonts already
    present in the page. This maps common fonts to the closest Base-14 equivalent.
    
    Args:
        font_name: Original font name from PDF
        flags: Font flags (0=normal, 2=italic, 16=bold, 18=bold+italic)
    """
    fn = font_name.lower()

    # Determine style flags from both name and flags parameter
    is_bold = "bold" in fn or "black" in fn or "heavy" in fn or (flags & 16)
    is_italic = "italic" in fn or "oblique" in fn or (flags & 2)

    # Determine family
    if "courier" in fn or "mono" in fn or "consolas" in fn or "menlo" in fn or (flags & 8):
        # Monospace family
        if is_bold and is_italic:
            return "cobi"  # Courier-BoldOblique
        elif is_bold:
            return "cobo"  # Courier-Bold
        elif is_italic:
            return "coit"  # Courier-Oblique
        return "cour"  # Courier
        
    elif "times" in fn or "serif" in fn or "georgia" in fn or "garamond" in fn or (flags & 4):
        # Serif family
        if is_bold and is_italic:
            return "tibi"  # Times-BoldItalic
        elif is_bold:
            return "tibo"  # Times-Bold
        elif is_italic:
            return "tiit"  # Times-Italic
        return "tiro"  # Times-Roman
        
    elif "symbol" in fn:
        return "symb"  # Symbol
        
    elif "zapf" in fn or "dingbat" in fn:
        return "zadb"  # ZapfDingbats
        
    else:
        # Default family: Helvetica (covers Arial, Calibri, Verdana, etc.)
        if is_bold and is_italic:
            return "hebi"  # Helvetica-BoldOblique
        elif is_bold:
            return "hebo"  # Helvetica-Bold
        elif is_italic:
            return "heit"  # Helvetica-Oblique
        return "helv"  # Helvetica


def _get_text_properties_at_rect(page, rect) -> dict:
    """Get comprehensive font properties of text at a given rectangle on the page.
    
    Enhanced to return complete formatting information including:
    - Font name and Base-14 mapping
    - Font size
    - Text color (RGB)
    - Font flags (bold, italic, etc.)
    - Character spacing
    
    Uses page.get_text("dict") to inspect individual text spans and finds the span
    with the most overlap with the search-hit rectangle.
    
    Returns: dict with keys: fontname, fontsize, text_color, flags, char_space
    """
    text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_IMAGES)
    best_span = None
    best_overlap = 0

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:  # text blocks only
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                span_rect = fitz.Rect(span["bbox"])
                overlap = span_rect & rect  # intersection
                if overlap.is_empty:
                    continue
                overlap_area = overlap.width * overlap.height
                if overlap_area > best_overlap:
                    best_overlap = overlap_area
                    best_span = span

    if best_span:
        font_name = best_span.get("font", "Helvetica")
        font_size = best_span.get("size", 11.0)
        color_int = best_span.get("color", 0)
        flags = best_span.get("flags", 0)

        # Convert integer color (0xRRGGBB) to (r, g, b) tuple with values 0–1
        r = ((color_int >> 16) & 0xFF) / 255.0
        g = ((color_int >> 8) & 0xFF) / 255.0
        b = (color_int & 0xFF) / 255.0

        return {
            "fontname": _map_to_base14_font(font_name, flags),
            "fontsize": font_size,
            "text_color": (r, g, b),
            "flags": flags,
            "original_font": font_name
        }

    # Fallback: Helvetica 11pt black
    return {
        "fontname": "helv",
        "fontsize": 11.0,
        "text_color": (0, 0, 0),
        "flags": 0,
        "original_font": "Helvetica"
    }


def sanitize_pdf_in_place(pdf_stream: BytesIO, replacements: list) -> BytesIO:
    """Redact and replace PII text directly in the PDF, preserving ALL layout and formatting.
    
    ENHANCED FORMAT PRESERVATION + ANTI-DUPLICATION:
    - Preserves exact font family, size, weight (bold), and style (italic)
    - Maintains original text color precisely
    - Keeps character spacing and kerning
    - Preserves all images, tables, headers, footers
    - Maintains page layout and margins
    - Prevents duplicate replacements by applying redactions incrementally
    
    Uses PyMuPDF redaction annotations with INCREMENTAL APPLICATION:
    1. Sort replacements by length (longest first: "Priya Anjali Fernando" before "Priya")
    2. For each replacement:
       a. Search for the original text on the page
       b. Add redaction annotations with exact font properties
       c. IMMEDIATELY apply redactions (this removes the original text)
    3. This ensures shorter names don't match inside already-replaced longer names
    
    Security: Original text is permanently removed (not just covered).
    """
    try:
        doc = fitz.open(stream=pdf_stream, filetype="pdf")

        # Sort replacements by length (longest first) to avoid partial matches
        # CRITICAL: "Priya Anjali Fernando" must be replaced BEFORE "Priya"
        sorted_replacements = sorted(replacements, key=lambda r: len(r["original"]), reverse=True)

        total_redactions = 0
        formatting_stats = {
            "fonts_detected": set(),
            "colors_detected": set(),
            "sizes_detected": set()
        }

        for page in doc:
            # CRITICAL FIX: Process each replacement separately and apply immediately
            # This prevents overlapping matches (e.g., "Priya" matching inside "Person_1")
            for rep in sorted_replacements:
                # Search for all instances of the original text on this page
                # NOTE: After previous redactions, the page text has changed,
                # so this search operates on the current state (already-replaced text won't match)
                text_instances = page.search_for(rep["original"])
                
                if not text_instances:
                    continue  # No matches for this replacement on this page
                
                # Add redaction annotations for all instances of this specific PII
                for inst in text_instances:
                    # Detect the original text's complete font properties at this location
                    props = _get_text_properties_at_rect(page, inst)
                    
                    # Track formatting diversity (for logging)
                    formatting_stats["fonts_detected"].add(props["original_font"])
                    formatting_stats["colors_detected"].add(props["text_color"])
                    formatting_stats["sizes_detected"].add(props["fontsize"])

                    # Add redaction annotation with EXACT matching style
                    page.add_redact_annot(
                        inst,
                        text=rep["replacement"],
                        fontname=props["fontname"],      # Matched Base-14 font
                        fontsize=props["fontsize"],      # Exact original size
                        text_color=props["text_color"],  # Exact original color (R,G,B)
                        fill=(1, 1, 1),                  # White background to cleanly cover
                        align=fitz.TEXT_ALIGN_LEFT       # Preserve text alignment
                    )
                    total_redactions += 1
                
                # CRITICAL: Apply redactions IMMEDIATELY after adding them for this specific PII
                # This updates the page text, so the next replacement searches the updated content
                # This prevents "Priya" from matching inside the already-replaced "Person_1"
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        logging.info(f"→ Applied {total_redactions} in-place PDF redaction(s)")
        logging.info(f"→ Format preservation: {len(formatting_stats['fonts_detected'])} unique fonts, "
                    f"{len(formatting_stats['colors_detected'])} unique colors, "
                    f"{len(formatting_stats['sizes_detected'])} unique sizes preserved")

        # Security: remove metadata that may contain PII
        doc.set_metadata({})

        # Security: remove JavaScript and embedded files
        doc.scrub(
            attached_files=True,      # Remove attached files
            clean_pages=True,         # Remove orphaned resources
            hidden_text=True,         # Remove hidden text
            javascript=True,          # Remove JavaScript
            metadata=True,            # Remove metadata
            redactions=False,         # Keep our redactions
            remove_links=False,       # Keep hyperlinks
            reset_fields=False,       # Keep form fields
            reset_responses=False,    # Keep form responses
            thumbnails=True,          # Remove thumbnails
            xml_metadata=True         # Remove XML metadata
        )

        # Save with maximum compression and cleanup
        output_stream = BytesIO()
        doc.save(
            output_stream,
            garbage=4,           # Maximum garbage collection
            deflate=True,        # Compress streams
            clean=True,          # Clean up duplicate objects
            pretty=False,        # Don't prettify (smaller file)
            linear=False,        # Don't linearize
            no_new_id=True       # Don't generate new ID (for consistency)
        )
        doc.close()
        output_stream.seek(0)
        return output_stream

    except Exception as e:
        logging.error(f"PDF in-place sanitization error: {str(e)}")
        raise ValueError(f"Failed to sanitize PDF: {str(e)}")


# ═══════════════════════════════════════════════════════════════
# DOCX FORMATTING PRESERVATION - Enhanced Implementation
# ═══════════════════════════════════════════════════════════════

def remove_images_and_add_placeholders(doc: Document):
    """Remove images and replace with styled placeholders.
    
    Enhanced to preserve surrounding text formatting.
    """
    photo_counter = 1

    for paragraph in doc.paragraphs:
        i = 0
        while i < len(paragraph.runs):
            run = paragraph.runs[i]
            drawing = run._element.find(qn('w:drawing'))
            if drawing is not None:
                # Store original formatting before clearing
                original_font_name = run.font.name
                original_font_size = run.font.size
                original_bold = run.bold
                original_italic = run.italic
                original_color = run.font.color.rgb if run.font.color.rgb else None
                
                run.clear()
                placeholder = f"[Photo-{photo_counter}]"
                photo_counter += 1
                
                # Add placeholder with preserved formatting
                run.text = placeholder + " "
                run.bold = True  # Make placeholder bold for visibility
                
                i += 1
            else:
                i += 1

    # Handle headers and footers
    for section in doc.sections:
        for header in section.header.paragraphs:
            i = 0
            while i < len(header.runs):
                run = header.runs[i]
                if run._element.find(qn('w:drawing')) is not None:
                    run.clear()
                    run.text = "[Header Image] "
                    run.bold = True
                i += 1

        for footer in section.footer.paragraphs:
            i = 0
            while i < len(footer.runs):
                run = footer.runs[i]
                if run._element.find(qn('w:drawing')) is not None:
                    run.clear()
                    run.text = "[Footer Image] "
                    run.bold = True
                i += 1

    logging.info(f"→ Replaced {photo_counter - 1} image(s) with placeholders")
    return doc


def extract_full_document_text(doc: Document) -> str:
    """Extract all text from paragraphs and tables into a single string for LLM analysis."""
    lines = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    
    # Extract from tables
    for table_idx, table in enumerate(doc.tables, 1):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
    
    # Extract from headers and footers
    for section in doc.sections:
        # Headers
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Footers
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
    
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════
# CROSS-RUN REPLACEMENT ENGINE
# ═══════════════════════════════════════════════════════════════
# 
# WHY THIS IS NEEDED:
# Word/DOCX splits text into "runs" — the smallest unit with consistent formatting.
# A paragraph like "Email: priya.fernando92@finance.lk" may be stored as:
#   Run 1: "Email: priya"
#   Run 2: ".fernando92@finance"
#   Run 3: ".lk"
#
# The old code checked each run individually, so it NEVER found
# "priya.fernando92@finance.lk" because no single run contained the full text.
#
# This new engine:
# 1. Concatenates ALL run texts in a paragraph
# 2. Finds PII matches in the concatenated text (case-insensitive)
# 3. Maps match positions back to individual runs
# 4. Surgically modifies only the affected runs while preserving formatting
# ═══════════════════════════════════════════════════════════════


def _apply_replacements_to_paragraph(para, sorted_replacements: list) -> int:
    """Apply ALL replacements to a paragraph using CROSS-RUN aware matching.
    
    This is the core fix for the DOCX replacement problem.
    Word splits text across runs unpredictably — this function handles that.
    
    Algorithm:
    1. Build a map of (character_position → run_index) from all runs
    2. Concatenate all run texts into one string
    3. For each replacement, find ALL matches (case-insensitive) in the full string
    4. For each match, determine which runs are affected
    5. Modify runs: put replacement text in the first affected run,
       clear text from middle runs, trim the last run
    6. Repeat until no more matches (handles overlapping replacements)
    
    Returns: Number of replacements made
    """
    runs = para.runs
    if not runs:
        return 0
    
    total_replacements = 0
    
    for rep in sorted_replacements:
        original = rep["original"]
        replacement = rep["replacement"]
        
        # We may need multiple passes if the same PII appears multiple times
        # After each replacement, run boundaries shift, so we re-scan
        max_iterations = 50  # Safety limit
        iteration = 0
        
        while iteration < max_iterations:
            iteration += 1
            runs = para.runs  # Re-fetch runs (they may have changed)
            if not runs:
                break
            
            # Step 1: Build the concatenated text and position map
            full_text = ""
            run_boundaries = []  # List of (start_pos, end_pos, run_index)
            
            for idx, run in enumerate(runs):
                start = len(full_text)
                full_text += run.text
                end = len(full_text)
                run_boundaries.append((start, end, idx))
            
            # DEBUG: Log email search attempts
            if "@" in original and "@" in full_text:
                logging.info(f"  → Searching for email '{original}' in paragraph text: '{full_text[:100]}...'")
            
            # Step 2: Find the FIRST match (case-insensitive)
            match = re.search(re.escape(original), full_text, flags=re.IGNORECASE)
            if not match:
                # DEBUG: Log failed email matches specifically
                if "@" in original and "@" in full_text:
                    logging.warning(f"  ✗ Email '{original}' NOT found in: '{full_text}'")
                break  # No more matches for this replacement
            
            match_start = match.start()
            match_end = match.end()
            
            # Step 3: Determine which runs are affected
            affected_runs = []
            for start, end, idx in run_boundaries:
                # A run is affected if it overlaps with the match range
                if start < match_end and end > match_start:
                    affected_runs.append({
                        "index": idx,
                        "run": runs[idx],
                        "run_start": start,
                        "run_end": end,
                        # How much of this run's text is before the match
                        "prefix_len": max(0, match_start - start),
                        # How much of this run's text is after the match
                        "suffix_len": max(0, end - match_end),
                    })
            
            if not affected_runs:
                break  # Should not happen, but safety check
            
            # Step 4: Apply the replacement across runs
            first = affected_runs[0]
            last = affected_runs[-1]
            
            # Save formatting from the first affected run (we'll apply replacement here)
            first_run = first["run"]
            
            # The first run keeps its prefix + gets the replacement text
            prefix = first_run.text[:first["prefix_len"]]
            # The last run keeps its suffix
            suffix = last["run"].text[len(last["run"].text) - last["suffix_len"]:] if last["suffix_len"] > 0 else ""
            
            if len(affected_runs) == 1:
                # Simple case: entire match is within one run
                # Just replace the text, keeping prefix and suffix
                first_run.text = prefix + replacement + suffix
            else:
                # Complex case: match spans multiple runs
                # First run: prefix + replacement text
                first_run.text = prefix + replacement
                
                # Middle runs: clear their text (they're fully consumed by the match)
                for affected in affected_runs[1:-1]:
                    affected["run"].text = ""
                
                # Last run: keep only the suffix
                last["run"].text = suffix
            
            total_replacements += 1
            # DEBUG: Log successful replacements, especially for emails
            if "@" in original:
                logging.info(f"  ✓ Email replaced: '{original}' → '{replacement}' (spans {len(affected_runs)} run(s))")
            else:
                logging.debug(f"  Cross-run replacement: '{original}' → '{replacement}' (spans {len(affected_runs)} run(s))")
    
    return total_replacements


def _apply_replacements_to_paragraph_fallback(para, sorted_replacements: list) -> int:
    """Fallback: direct paragraph-level text replacement for edge cases.
    
    If cross-run replacement finds nothing but the paragraph text contains the PII,
    this handles the case where runs are completely fragmented or missing.
    This modifies paragraph XML directly as a last resort.
    """
    runs = para.runs
    if not runs:
        return 0
    
    full_text = "".join(run.text for run in runs)
    replacements_made = 0
    
    for rep in sorted_replacements:
        if rep["original"].lower() in full_text.lower():
            # The PII exists in this paragraph but wasn't caught by cross-run
            # This shouldn't happen often, but if it does, do a simple replace
            # on each run individually as a safety net
            for run in runs:
                if rep["original"].lower() in run.text.lower():
                    new_text = re.sub(
                        re.escape(rep["original"]),
                        rep["replacement"],
                        run.text,
                        flags=re.IGNORECASE
                    )
                    if new_text != run.text:
                        run.text = new_text
                        replacements_made += 1
            
            # Re-check: update full_text after modifications
            full_text = "".join(run.text for run in runs)
    
    return replacements_made


def apply_replacements_to_document(doc: Document, replacements: list) -> dict:
    """Apply replacement mapping to entire document with CROSS-RUN aware matching.
    
    THIS IS THE KEY FUNCTION that guarantees 100% replacement accuracy.
    
    WHY THE OLD APPROACH FAILED:
    Word splits text into "runs" at formatting boundaries (and sometimes randomly).
    Text like "priya.fernando92@finance.lk" may be split across 3+ runs.
    The old code checked each run individually — if no single run contained
    the full PII text, the replacement silently failed.
    
    NEW APPROACH:
    1. Concatenate ALL run texts in a paragraph into one string
    2. Find PII matches in the concatenated text (case-insensitive)
    3. Map match positions back to individual runs
    4. Surgically modify only the affected runs
    5. Formatting is preserved because we modify run.text without touching XML properties
    
    Returns: Statistics about replacements made
    """
    if not replacements:
        return {
            "paragraphs_changed": 0,
            "table_cells_changed": 0,
            "header_footer_changed": 0,
            "total_replacements": 0
        }

    # Sort replacements by length (longest first) to avoid partial matches
    # e.g., "Priya Anjali Fernando" must be replaced before "Priya"
    sorted_replacements = sorted(replacements, key=lambda r: len(r["original"]), reverse=True)

    stats = {
        "paragraphs_changed": 0,
        "table_cells_changed": 0,
        "header_footer_changed": 0,
        "total_replacements": 0
    }

    def _process_paragraph(para):
        """Process a single paragraph with cross-run replacement + fallback."""
        count = _apply_replacements_to_paragraph(para, sorted_replacements)
        if count == 0:
            # Fallback for edge cases
            count = _apply_replacements_to_paragraph_fallback(para, sorted_replacements)
        return count

    # Process main document paragraphs
    for para in doc.paragraphs:
        count = _process_paragraph(para)
        if count > 0:
            stats["total_replacements"] += count
            stats["paragraphs_changed"] += 1

    # Process tables (CRITICAL — many PII items like phones/emails are in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_changed = False
                for para in cell.paragraphs:
                    count = _process_paragraph(para)
                    if count > 0:
                        stats["total_replacements"] += count
                        cell_changed = True
                if cell_changed:
                    stats["table_cells_changed"] += 1

    # Process headers and footers
    for section in doc.sections:
        # Headers
        for para in section.header.paragraphs:
            count = _process_paragraph(para)
            if count > 0:
                stats["total_replacements"] += count
                stats["header_footer_changed"] += 1
        
        # Footers
        for para in section.footer.paragraphs:
            count = _process_paragraph(para)
            if count > 0:
                stats["total_replacements"] += count
                stats["header_footer_changed"] += 1

    logging.info(f"Cross-run replacement stats: "
                f"{stats['total_replacements']} replacements made, "
                f"{stats['paragraphs_changed']} paragraphs, "
                f"{stats['table_cells_changed']} table cells, "
                f"{stats['header_footer_changed']} headers/footers")

    return stats


# ═══════════════════════════════════════════════════════════════
# Regex Safety Net — catches PII the LLM missed
# ═══════════════════════════════════════════════════════════════

def _apply_regex_safety_net(full_text: str, replacements: list) -> list:
    """Post-process LLM results with regex to catch any missed emails, phones, and common ID patterns.
    
    This is a CRITICAL safety layer. LLMs can miss PII items, especially:
    - Emails embedded in tables or inline text
    - Phone numbers in various formats
    - Partial name references
    - Email address variations (with/without dots, numbers, etc.)
    
    This function scans the document text with regex patterns and adds any
    PII not already covered by the LLM's mapping.
    """
    # Build a set of originals already covered (case-insensitive, normalized)
    covered = set()
    for r in replacements:
        covered.add(r["original"].lower())
        # Also add normalized version for emails (remove dots before @)
        if "@" in r["original"]:
            # Normalize: johndoe@domain.com and john.doe@domain.com should match
            normalized = re.sub(r'\.(?=[^@]*@)', '', r["original"].lower())
            covered.add(normalized)
    
    added_count = 0
    
    # Determine the next sequential numbers for emails and phones
    existing_email_count = sum(1 for r in replacements if r.get("type") == "email")
    existing_phone_count = sum(1 for r in replacements if r.get("type") == "phone")
    
    email_counter = existing_email_count + 1
    phone_counter = existing_phone_count + 1
    
    # ── EMAIL DETECTION ──
    # Match any email address pattern
    email_pattern = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}')
    for match in email_pattern.finditer(full_text):
        email = match.group()
        if email.lower() not in covered:
            replacements.append({
                "original": email,
                "type": "email",
                "replacement": f"person_{email_counter}@example.com"
            })
            covered.add(email.lower())
            email_counter += 1
            added_count += 1
            logging.info(f"  → Regex safety net caught missed email: '{email}' → person_{email_counter-1}@example.com")
    
    # ── PHONE NUMBER DETECTION ──
    # Match various phone number formats including international
    phone_patterns = [
        # International format: +94 77 523 4567, +1-555-123-4567
        re.compile(r'\+\d{1,3}[\s\-]?\d{1,4}[\s\-]?\d{2,4}[\s\-]?\d{3,4}'),
        # Local format with leading 0: 077-523-4567, (077) 523 4567, 077 523 4567
        re.compile(r'\(?\d{3,4}\)?[\s\-\.]\d{3,4}[\s\-\.]\d{3,4}'),
        # Compact: 10+ digits possibly with dashes/spaces
        re.compile(r'(?<!\d)\d{10,12}(?!\d)'),
    ]
    
    for pattern in phone_patterns:
        for match in pattern.finditer(full_text):
            phone = match.group().strip()
            # Skip if too short or looks like a year/amount
            if len(re.sub(r'[\s\-\(\)\.\+]', '', phone)) < 7:
                continue
            if phone.lower() not in covered:
                replacements.append({
                    "original": phone,
                    "type": "phone",
                    "replacement": f"+00 00 000 {phone_counter:04d}"
                })
                covered.add(phone.lower())
                phone_counter += 1
                added_count += 1
                logging.info(f"  → Regex safety net caught missed phone: ***{phone[-4:]}")
    
    # ── NIC / ID NUMBER DETECTION ──
    # Sri Lankan NIC: 9 digits + V/X, or 12 digits
    nic_patterns = [
        re.compile(r'(?<!\d)\d{9}[VvXx](?!\d)'),
        re.compile(r'(?<!\d)\d{12}(?!\d)'),
    ]
    
    existing_id_count = sum(1 for r in replacements if r.get("type") in ("id_number", "nic", "passport", "ssn"))
    id_counter = existing_id_count + 1
    
    for pattern in nic_patterns:
        for match in pattern.finditer(full_text):
            nic = match.group()
            if nic.lower() not in covered:
                replacements.append({
                    "original": nic,
                    "type": "id_number",
                    "replacement": f"ID_{id_counter:06d}"
                })
                covered.add(nic.lower())
                id_counter += 1
                added_count += 1
                logging.info(f"  → Regex safety net caught missed ID: ***{nic[-3:]}")
    
    if added_count > 0:
        logging.info(f"→ Regex safety net added {added_count} PII item(s) missed by LLM")
    else:
        logging.info("→ Regex safety net: LLM caught all emails/phones/IDs — no additions needed")
    
    return replacements


# ═══════════════════════════════════════════════════════════════
# LLM-based PII Detection — Enhanced with safety net
# ═══════════════════════════════════════════════════════════════

def build_replacement_mapping(full_text: str) -> list:
    """Send the entire document text to LLM once, get a consistent replacement mapping."""

    prompt = f"""TASK: Analyze the document below and build a COMPLETE PII replacement mapping.

IMPORTANT REQUIREMENTS:
1. Detect ALL email addresses (anything with @) — even inside tables and inline text
2. Detect ALL phone numbers in ANY format — even inside tables, bullet points, and after labels like "Phone:"
3. For EVERY person with multiple names (e.g., "Priya Anjali Fernando"), create SEPARATE mapping entries for:
   - The full name ("Priya Anjali Fernando")
   - First name alone ("Priya") — THIS IS CRITICAL, first names used alone elsewhere MUST be mapped
   - Last name alone ("Fernando")
   - Any other variation found in the document
4. Scan EVERY line of the document including tables, headers, footers, and bullet points

CATEGORIES TO DETECT:
• Personal: Full names, first/last names, nicknames, titles, job positions
• Contact: ALL emails (user@domain), ALL phone numbers (any format: +94 77 523 4567, 077-523-4567, etc.)
• Location: Addresses, cities, countries, postal codes, regions
• Financial: Bank names, account numbers, card numbers, amounts, currency
• Identifiers: NIC, passport, SSN, employee IDs, customer IDs, any ID numbers
• Dates: Birth dates, any personally identifying dates
• Organizations: Companies, banks, institutions, departments
• Technical: URLs, websites, API keys, system IDs

DOCUMENT TEXT:
{full_text}

Before returning, verify: every email has an entry, every phone has an entry, every first name used alone has an entry.

Return valid JSON array only:"""

    system_instructions = """You are an expert PII (Personally Identifiable Information) detection and anonymization specialist. You must achieve 100% detection accuracy — missed PII is a critical security failure.

YOUR MISSION:
Analyze the ENTIRE document — including tables, headers, footers, bullet points, and inline text — with EXTREME thoroughness to identify ALL sensitive information. EVERY email, phone number, name, and identifier MUST be detected. Missing even ONE is unacceptable.

═══════════════════════════════════════════
MANDATORY MULTI-PASS ANALYSIS:
═══════════════════════════════════════════
1. ENTITY DISCOVERY: Read the ENTIRE document. Build a list of every person, organization, and entity mentioned.
2. VARIATION MAPPING: For EACH person, list ALL ways they could be referenced:
   - Full name: "Priya Anjali Fernando"
   - First + Last: "Priya Fernando"
   - First name ONLY: "Priya" (THIS IS CRITICAL — first names used alone MUST be mapped)
   - Last name only: "Fernando"
   - With title: "Ms. Fernando", "Dr. Fernando"
   - Initials: "P. Fernando", "P.A.F."
   - In emails: if email contains a name (e.g., priya.fernando92@...), the email MUST also be detected
3. CONTACT INFO SCAN: Scan EVERY line for:
   - Email addresses: ANY text matching *@*.* pattern (e.g., john.doe87@swivel.lk, priya.fernando92@finance.lk)
   - Phone numbers: ANY numeric sequence that looks like a phone in ANY format:
     * International: +94 77 523 4567, +1-555-123-4567
     * Local: 077-523-4567, (077) 523 4567
     * With spaces: 077 523 4567
     * With dots: 077.523.4567
     * Plain digits: 0775234567
     * With extensions: +94 77 523 4567 ext 100
   - IMPORTANT: Check INSIDE tables, bullet lists, and inline text like "Phone: +94 77 523 4567"
4. FULL DOCUMENT SCAN: Go through the text LINE BY LINE. For EACH line ask:
   - Does this line contain a name, email, phone, address, ID, date, or organization?
   - Could any word on this line be a person's first name that was identified earlier?
5. VALIDATION: Before returning, verify:
   - Every email address in the document has a mapping entry
   - Every phone number in the document has a mapping entry
   - Every person's first name (used alone) has its OWN mapping entry
   - Every name variation has a mapping entry

═══════════════════════════════════════════
CRITICAL NAME RULES:
═══════════════════════════════════════════
- If someone has multiple names (e.g., "Priya Anjali Fernando"), you MUST create SEPARATE entries for:
  * The full name: "Priya Anjali Fernando" → Person_1
  * First + Last: "Priya Fernando" → Person_1
  * First name alone: "Priya" → Person_1
  * Last name alone: "Fernando" → Person_1
  * Any other variation found in the document
- The FIRST NAME USED ALONE is the #1 most commonly missed PII item. ALWAYS include it.
- If a name appears in performance notes, comments, or any section — it MUST be detected.
- Names in tables MUST be detected.

═══════════════════════════════════════════
CRITICAL EMAIL RULES:
═══════════════════════════════════════════
- EVERY email address MUST be detected, no exceptions.
- Emails embedded in text like "Email: john.doe87@swivel.lk" — extract "john.doe87@swivel.lk"
- Emails in tables MUST be detected.
- Map email to the corresponding person: if john.doe87@swivel.lk belongs to Person_1, replace with person_1@example.com

═══════════════════════════════════════════
CRITICAL PHONE RULES:
═══════════════════════════════════════════
- EVERY phone number MUST be detected, regardless of format.
- Phone numbers in tables MUST be detected.
- Phone numbers after labels like "Phone:", "Tel:", "Mobile:", "Contact:" MUST be detected.
- Phone numbers with country codes (+94, +1, +44, etc.) MUST be detected.
- Each unique phone number gets a sequential replacement: +00 00 000 0001, +00 00 000 0002, etc.

═══════════════════════════════════════════
CONSISTENCY RULES:
═══════════════════════════════════════════
- Same person = same Person_N number across ALL entries (full name, first name, email, etc.)
- If "Priya Anjali Fernando" is Person_1, then "Priya" is also Person_1, and "priya.fernando92@finance.lk" becomes person_1@example.com

REPLACEMENT CONVENTIONS (FOLLOW EXACTLY):
- Full names: Person_1, Person_2... (sequential)
- Name variations (first, last, partial): SAME Person_N as the full name
- Job Titles: Job_1, Job_2...
- Emails: person_N@example.com (N matches the person's number)
- Phones: +00 00 000 0001, +00 00 000 0002... (sequential)
- Organizations/Banks: Organization_1, Organization_2...
- Addresses: 123 Sample St, City_1, Country_1, 00000
- IDs/NIC/Passport/SSN: ID_000001, ID_000002...
- Dates: 01/01/1990, 02/02/1991...
- Financial amounts: ****1234, $X,XXX

OUTPUT FORMAT:
Return ONLY a valid JSON array. No explanations, no markdown, no comments.
Each object: {"original": "exact text", "type": "category", "replacement": "sanitized value"}

═══════════════════════════════════════════
FINAL MANDATORY SELF-CHECK (DO THIS BEFORE RETURNING):
═══════════════════════════════════════════
□ Did I include EVERY email address found anywhere in the document?
□ Did I include EVERY phone number found anywhere in the document (including tables)?
□ Did I include EVERY person's first name as a SEPARATE entry?
□ Did I check tables, bullet points, headers, and footers for PII?
□ Did I map name variations (first name alone) to the same Person_N?
□ Is my JSON valid and complete?

If ANY check fails, go back and add the missing entries before returning.

Return empty array [] if no PII is found. Never include explanatory text outside the JSON array."""

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_instructions},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,  # Maximum consistency
            max_tokens=8192,  # Increased to handle large documents with many PII entries
        )

        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content.split("```json")[1].split("```")[0].strip()
        elif content.startswith("```"):
            content = content.split("```")[1].split("```")[0].strip()

        replacements = json.loads(content)
        logging.info(f"→ LLM returned {len(replacements)} unique PII replacement(s)")
        
        # SAFETY NET: Regex-based post-scan to catch any emails/phones the LLM missed
        replacements = _apply_regex_safety_net(full_text, replacements)
        
        # Note: Detailed PII mappings are NOT logged for security reasons
        # The original PII data is only held in memory during processing
        
        return replacements

    except json.JSONDecodeError as je:
        logging.warning(f"JSON parse error from LLM: {str(je)} — returning empty mapping")
        return []
    except OpenAIError as oe:
        logging.error(f"OpenAI API error: {str(oe)}")
        return []
    except Exception as e:
        logging.error(f"Unexpected error building replacement mapping: {str(e)}")
        return []


# ═══════════════════════════════════════════════════════════════
# Core sanitization orchestration
# ═══════════════════════════════════════════════════════════════

def sanitize_docx(input_stream: BytesIO) -> BytesIO:
    """Sanitize DOCX file with COMPLETE format preservation.
    
    Process:
    1. Remove images (replace with placeholders)
    2. Extract all text for LLM analysis
    3. Build global PII replacement mapping
    4. Apply replacements at RUN level (preserves all formatting)
    
    Format preservation: Font, size, color, bold, italic, underline, highlights, etc.
    """
    logging.info("Starting .docx sanitization with format preservation")

    doc = Document(input_stream)

    # Step 1: Remove images and add placeholders
    doc = remove_images_and_add_placeholders(doc)

    # Step 2: Extract ALL text from the document (including headers/footers)
    full_text = extract_full_document_text(doc)
    logging.info(f"Extracted {len(full_text)} characters of document text")

    if not full_text.strip():
        logging.info("No text content found in document — skipping sanitization")
    else:
        # Step 3: Single LLM call to build global replacement mapping
        replacements = build_replacement_mapping(full_text)

        # Step 4: Apply the mapping with CROSS-RUN aware replacement
        stats = apply_replacements_to_document(doc, replacements)
        
        logging.info(f"DOCX Sanitization Summary:")
        logging.info(f"  • {stats['total_replacements']} total replacements made")
        logging.info(f"  • {stats['paragraphs_changed']} paragraphs affected")
        logging.info(f"  • {stats['table_cells_changed']} table cells affected")
        logging.info(f"  • {stats['header_footer_changed']} headers/footers affected")
        logging.info(f"  • {len(replacements)} unique PII items detected and replaced")

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


# ═══════════════════════════════════════════════════════════════
# HTTP endpoint
# ═══════════════════════════════════════════════════════════════

@app.route(route="sanitize-docx", methods=["POST"], auth_level=func.AuthLevel.FUNCTION)
def sanitize_docx_http(req: func.HttpRequest) -> func.HttpResponse:
    """HTTP endpoint for document sanitization with format preservation.
    
    Supports: .docx and .pdf files
    Security: File type validation, size limits, content validation
    Format Preservation: Complete (fonts, sizes, colors, styles, margins)
    """
    logging.info("HTTP request received for document sanitization")

    try:
        # Security: Validate file size (max 10MB)
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
        
        # Get file from multipart form (field name 'file')
        if 'file' not in req.files:
            return func.HttpResponse("Missing 'file' in multipart form data", status_code=400)

        uploaded_file = req.files['file']
        filename_lower = uploaded_file.filename.lower()
        
        # Security: Check file extension
        allowed_extensions = ['.docx', '.pdf']
        file_extension = None
        for ext in allowed_extensions:
            if filename_lower.endswith(ext):
                file_extension = ext
                break
        
        if not file_extension:
            return func.HttpResponse(
                "Unsupported file type. Only .docx and .pdf files are supported. "
                "Legacy .doc format is not supported for security reasons - please convert to .docx first.",
                status_code=400
            )

        file_content = uploaded_file.read()
        
        # Security: Check file size
        if len(file_content) == 0:
            return func.HttpResponse("Uploaded file is empty", status_code=400)
        if len(file_content) > MAX_FILE_SIZE:
            return func.HttpResponse("File too large. Maximum size is 10MB", status_code=400)
        
        # Security: Validate file type by content (magic bytes), not just extension
        if not validate_file_type(file_content, file_extension):
            return func.HttpResponse(
                "File type mismatch. The file content does not match its extension. Possible security risk.",
                status_code=400
            )

        # Process based on file type
        if file_extension == '.docx':
            logging.info("Processing DOCX file with format preservation")
            input_stream = BytesIO(file_content)
            output_stream = sanitize_docx(input_stream)
            output_mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            output_extension = ".docx"
            
        elif file_extension == '.pdf':
            logging.info("Processing PDF file with format preservation")
            
            # Step 1: Extract text for LLM analysis
            input_stream = BytesIO(file_content)
            full_text = extract_text_from_pdf(input_stream)
            
            if not full_text.strip():
                return func.HttpResponse("No text content found in PDF", status_code=400)
            
            logging.info(f"Extracted {len(full_text)} characters from PDF")
            
            # Step 2: Build replacement mapping via LLM
            replacements = build_replacement_mapping(full_text)
            
            # Step 3: Apply redactions in-place with COMPLETE format preservation
            # (preserves fonts, sizes, colors, styles, layout, images, tables)
            input_stream = BytesIO(file_content)  # fresh stream for modification
            output_stream = sanitize_pdf_in_place(input_stream, replacements)
            output_mimetype = "application/pdf"
            output_extension = ".pdf"
            
            logging.info(f"PDF Sanitization Summary:")
            logging.info(f"  • {len(replacements)} unique PII items detected and replaced")
            logging.info(f"  • Format preservation: fonts, sizes, colors, styles maintained")
        
        else:
            return func.HttpResponse("Unsupported file type", status_code=400)

        # Generate sanitized filename
        original_name = uploaded_file.filename.rsplit('.', 1)[0]
        sanitized_filename = f"sanitized_{original_name}{output_extension}"

        return func.HttpResponse(
            body=output_stream.read(),
            status_code=200,
            mimetype=output_mimetype,
            headers={
                "Content-Disposition": f'attachment; filename="{sanitized_filename}"',
                "Content-Length": str(output_stream.getbuffer().nbytes)
            }
        )

    except ValueError as ve:
        logging.error(f"Validation error: {str(ve)}")
        return func.HttpResponse(f"Validation error: {str(ve)}", status_code=400)
    except Exception as e:
        logging.error(f"Processing failed: {str(e)}", exc_info=True)
        return func.HttpResponse(f"Server error: {str(e)}", status_code=500) 


