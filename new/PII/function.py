import logging
import re
import uuid
from io import BytesIO
from pathlib import Path
from typing import Any, Final

import fitz  # PyMuPDF
from docx import Document
from docx.oxml.ns import qn


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """
    Extract plain text from a PDF given its raw bytes, using PyMuPDF,
    with robust error handling. This is suitable for HTTP uploads.

    Args:
        pdf_bytes: Raw PDF file contents.

    Returns:
        The concatenated text content of all pages.

    Raises:
        ValueError: If text extraction fails for any reason.
    """
    text_parts: list[str] = []

    try:
        pdf_stream = BytesIO(pdf_bytes)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")
        for page_num, page in enumerate(doc, start=1):
            try:
                text = page.get_text()
                if text and text.strip():
                    text_parts.append(text.rstrip())
            except Exception as e:  # per-page safeguard
                logging.warning(f"Could not extract text from page {page_num}: {e}")
                continue
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logging.error(f"PDF text extraction error from bytes: {e}")
        raise ValueError("Failed to extract text from PDF") from e


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract plain text from a PDF using PyMuPDF, with robust error handling.

    Args:
        pdf_path: Path to the PDF file.

    Returns:
        The concatenated text content of all pages.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If text extraction fails for any reason.
    """
    path = Path(pdf_path)
    if not path.is_file():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    try:
        pdf_bytes = path.read_bytes()
    except Exception as e:
        logging.error(f"Failed to read PDF file '{pdf_path}': {e}")
        raise

    return extract_text_from_pdf_bytes(pdf_bytes)


def sanitize_pdf_bytes(pdf_bytes: bytes, replacements: list[dict[str, Any]]) -> bytes:
    """
    Replace PII in a PDF with sanitized values in-place, preserving layout.
    Uses PyMuPDF redaction: search for each original text, redact and replace with
    the replacement value. Replacements are applied longest-first to avoid partial
    matches (e.g. "John M. Doe" before "John").

    Args:
        pdf_bytes: Raw PDF file contents.
        replacements: List of dicts with "original" and "replacement" keys
                      (e.g. [{"original": "John Doe", "type": "person", "replacement": "Person_1"}]).

    Returns:
        PDF file contents with replacements applied.

    Raises:
        ValueError: If PDF cannot be opened or processed.
    """
    if not replacements:
        return pdf_bytes

    # Normalize: ensure we have original/replacement; sort by len(original) descending
    repl_list = []
    for r in replacements:
        orig = r.get("original")
        repl = r.get("replacement")
        if orig is not None and repl is not None:
            repl_list.append({"original": str(orig), "replacement": str(repl)})
    repl_list.sort(key=lambda x: len(x["original"]), reverse=True)

    try:
        doc = fitz.open(stream=BytesIO(pdf_bytes), filetype="pdf")
    except Exception as e:
        logging.error(f"sanitize_pdf_bytes: failed to open PDF: {e}")
        raise ValueError("Failed to open PDF") from e

    for page in doc:
        for rep in repl_list:
            original = rep["original"]
            replacement = rep["replacement"]
            if _needs_word_boundary(original):
                # Short all-caps tokens (e.g. "IT", "HR"): filter search
                # results to only case-sensitive matches, preventing
                # "it" inside "within" / "City" from being redacted.
                all_instances = page.search_for(original)
                instances = []
                for rect in all_instances:
                    clip_text = page.get_text("text", clip=rect).strip()
                    if original in clip_text:  # case-sensitive check
                        instances.append(rect)
            else:
                instances = page.search_for(original)
            for rect in instances:
                try:
                    page.add_redact_annot(
                        rect,
                        text=replacement,
                        fontname="helv",
                        fontsize=11,
                        fill=(1, 1, 1),
                    )
                except Exception as e:
                    logging.warning(f"add_redact_annot failed for '{original[:30]}...': {e}")
            if instances:
                try:
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                except Exception as e:
                    logging.warning(f"apply_redactions failed: {e}")

    out = BytesIO()
    doc.save(out, deflate=True)
    doc.close()
    return out.getvalue()


def _paragraph_to_text(para) -> str:
    """
    Convert a python-docx paragraph to plain text, adding simple markers for lists.
    """
    raw = para.text.strip()
    if not raw:
        return ""

    # Best-effort detection of bullet / numbered list styles
    style_name = getattr(getattr(para, "style", None), "name", "") or ""
    is_list = any(keyword in style_name for keyword in ("List Bullet", "List Number", "Bullet", "Numbered"))

    return f"- {raw}" if is_list else raw


def _extract_full_docx_text(document: Document) -> str:
    """
    Extract text from a DOCX document, including:
    - body paragraphs
    - table cells
    - headers and footers
    """
    lines: list[str] = []

    # Body paragraphs
    for para in document.paragraphs:
        text = _paragraph_to_text(para)
        if text:
            lines.append(text)

    # Tables (rows × cells, including paragraphs inside cells)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = _paragraph_to_text(para)
                    if text:
                        lines.append(text)

    # Headers and footers
    for section in document.sections:
        for para in section.header.paragraphs:
            text = _paragraph_to_text(para)
            if text:
                lines.append(text)
        for para in section.footer.paragraphs:
            text = _paragraph_to_text(para)
            if text:
                lines.append(text)

    return "\n".join(lines)


def _needs_word_boundary(original: str) -> bool:
    """Check if a replacement original needs word-boundary matching.

    Short (<=3 chars) all-uppercase alphabetic originals like "IT", "HR"
    need word-boundary matching to avoid replacing substrings inside
    normal words (e.g. "it" in "within", "City", "audit").
    """
    return len(original) <= 3 and original.isupper() and original.isalpha()


def _build_search_pattern(original: str) -> tuple[str, int]:
    """Build regex pattern and flags for matching a replacement original.

    For short all-caps originals (e.g. "IT", "HR", "SAP"):
      - Word-boundary anchors (\\b) prevent substring matches
      - Case-sensitive matching ensures "IT" != "it"

    For everything else:
      - Plain escaped pattern with case-insensitive matching

    Returns:
        (pattern, re_flags)
    """
    escaped = re.escape(original)
    if _needs_word_boundary(original):
        return rf'\b{escaped}\b', 0
    return escaped, re.IGNORECASE


def _unwrap_hyperlinks(para) -> None:
    """Move runs out of w:hyperlink wrappers into the paragraph directly.

    python-docx's para.runs only returns w:r elements that are direct
    children of w:p.  Runs inside w:hyperlink are invisible and won't be
    processed by the replacement engine.  This function "unwraps" every
    hyperlink: its child w:r elements are promoted to direct children of
    the paragraph, and the w:hyperlink element is removed.

    Side-effects:
      - Hyperlink display text becomes normal paragraph text (still replaceable).
      - The hyperlink relationship (URL/mailto that may contain PII) is removed.
    """
    for hyperlink in para._element.findall(qn('w:hyperlink')):
        parent = hyperlink.getparent()
        idx = list(parent).index(hyperlink)
        for i, run_elem in enumerate(hyperlink.findall(qn('w:r'))):
            parent.insert(idx + i, run_elem)
        parent.remove(hyperlink)


def _apply_replacements_to_paragraph(para, sorted_replacements: list[dict[str, Any]]) -> int:
    """
    Apply replacements to a single paragraph with cross-run awareness.

    Handles three edge cases that cause missed or incorrect replacements:

    1. Cross-run text: Word splits text across runs; we concatenate run
       texts, find matches, and apply replacements across affected runs.
    2. Hyperlinks: Runs inside w:hyperlink elements are invisible to
       para.runs. We unwrap hyperlinks first so their text is replaceable
       and any PII-containing mailto/href is removed.
    3. Short all-caps tokens: Originals like "IT" or "HR" use word-boundary
       + case-sensitive matching to avoid replacing substrings (e.g. "it"
       inside "within").

    Longest-first order is assumed in sorted_replacements.
    """
    _unwrap_hyperlinks(para)
    runs = para.runs
    if not runs:
        return 0
    count = 0
    for rep in sorted_replacements:
        original = rep["original"]
        replacement = rep["replacement"]
        pattern, flags = _build_search_pattern(original)
        max_iter = 50
        while max_iter > 0:
            max_iter -= 1
            runs = para.runs
            if not runs:
                break
            full_text = ""
            run_boundaries = []
            for idx, run in enumerate(runs):
                start = len(full_text)
                full_text += run.text
                run_boundaries.append((start, len(full_text), idx))
            match = re.search(pattern, full_text, flags=flags)
            if not match:
                break
            match_start, match_end = match.start(), match.end()
            affected = []
            for start, end, idx in run_boundaries:
                if start < match_end and end > match_start:
                    prefix_len = max(0, match_start - start)
                    suffix_len = max(0, end - match_end)
                    affected.append({
                        "idx": idx, "run": runs[idx],
                        "prefix_len": prefix_len, "suffix_len": suffix_len,
                    })
            if not affected:
                break
            first, last = affected[0], affected[-1]
            prefix = first["run"].text[: first["prefix_len"]]
            suffix = (last["run"].text[len(last["run"].text) - last["suffix_len"] :]
                      if last["suffix_len"] > 0 else "")
            if len(affected) == 1:
                first["run"].text = prefix + replacement + suffix
            else:
                first["run"].text = prefix + replacement
                for a in affected[1:-1]:
                    a["run"].text = ""
                last["run"].text = suffix
            count += 1
    return count


def sanitize_docx_bytes(docx_bytes: bytes, replacements: list[dict[str, Any]]) -> bytes:
    """
    Replace PII in a DOCX with sanitized values. Processes body paragraphs,
    table cells, headers and footers. Uses cross-run replacement so text
    split across Word runs (e.g. "ruwan.silva@ops.lk" in multiple runs) is
    still replaced correctly. Replacements are applied longest-first.

    Args:
        docx_bytes: Raw DOCX file contents.
        replacements: List of dicts with "original" and "replacement" keys.

    Returns:
        DOCX file contents with replacements applied.

    Raises:
        ValueError: If DOCX cannot be opened or processed.
    """
    if not replacements:
        return docx_bytes

    repl_list = []
    for r in replacements:
        orig = r.get("original")
        repl = r.get("replacement")
        if orig is not None and repl is not None:
            repl_list.append({"original": str(orig), "replacement": str(repl)})
    repl_list.sort(key=lambda x: len(x["original"]), reverse=True)

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as e:
        logging.error(f"sanitize_docx_bytes: failed to open DOCX: {e}")
        raise ValueError("Failed to open DOCX") from e

    def process_para(paragraph):
        _apply_replacements_to_paragraph(paragraph, repl_list)

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            process_para(para)
        for para in section.footer.paragraphs:
            process_para(para)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract all visible text from a DOCX file given its raw bytes, including:
    - regular paragraphs
    - tables
    - headers and footers
    - bullet and numbered list paragraphs (marked with "- " prefix)
    """
    try:
        document = Document(BytesIO(docx_bytes))
    except Exception as e:
        logging.error(f"DOCX open error from bytes: {e}")
        raise ValueError("Failed to open DOCX document") from e

    return _extract_full_docx_text(document)


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract all visible text from a DOCX file, including:
    - regular paragraphs
    - tables
    - headers and footers
    - bullet and numbered list paragraphs (marked with "- " prefix)
    """
    path = Path(docx_path)
    if not path.is_file():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        docx_bytes = path.read_bytes()
    except Exception as e:
        logging.error(f"Failed to read DOCX file '{docx_path}': {e}")
        raise

    return extract_text_from_docx_bytes(docx_bytes)


SUPPORTED_EXTENSIONS: Final[set[str]] = {".pdf", ".docx"}

# MIME types for sanitized file responses
PDF_MIMETYPE: Final[str] = "application/pdf"
DOCX_MIMETYPE: Final[str] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def sanitize_document(
    file_bytes: bytes,
    filename: str,
    replacements: list[dict[str, Any]],
) -> tuple[bytes, str]:
    """
    Replace PII in an uploaded PDF or DOCX with the given replacements and
    return the modified file in its original format.

    Args:
        file_bytes: Raw file contents.
        filename: Original filename (used to determine type from extension).
        replacements: List of {"original": str, "replacement": str, ...}.

    Returns:
        (output_file_bytes, mimetype) for the sanitized document.

    Raises:
        ValueError: If file type is not supported or processing fails.
    """
    path = Path(filename)
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if extension == ".pdf":
        out_bytes = sanitize_pdf_bytes(file_bytes, replacements)
        return (out_bytes, PDF_MIMETYPE)
    else:
        out_bytes = sanitize_docx_bytes(file_bytes, replacements)
        return (out_bytes, DOCX_MIMETYPE)


def convert_to_text_file(input_path: str, output_dir: str) -> Path:
    """
    Read a PDF or DOCX file, extract its text, and save it as a .txt file
    with a unique identifier in the filename.

    The function chooses the appropriate extractor based on the file extension.

    Args:
        input_path: Path to the source PDF or DOCX file.
        output_dir: Directory where the generated text file will be stored.

    Returns:
        The path to the generated text file.

    Raises:
        FileNotFoundError: If the input file does not exist.
        ValueError: If the file extension is not supported.
    """
    source = Path(input_path)
    if not source.is_file():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    extension = source.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {extension}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")

    if extension == ".pdf":
        extracted_text = extract_text_from_pdf(str(source))
    else:  # ".docx"
        extracted_text = extract_text_from_docx(str(source))

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    unique_id = uuid.uuid4().hex
    output_file = output_path / f"{source.stem}_{unique_id}.txt"

    # Use UTF-8 to handle most text content robustly
    output_file.write_text(extracted_text, encoding="utf-8")

    return output_file

if __name__ == "__main__":
    # Example 1: Extract text and save as .txt
    # output = convert_to_text_file("input.pdf", ".")
    # print(f"Saved to: {output}")

    # Example 2: Sanitize a document with PII replacements (test locally)
    import sys
    if len(sys.argv) >= 3:
        # Usage: python function.py <input_file> <output_file> [replacements as JSON file path]
        # Or:    python function.py input.docx sanitized.docx
        # Replacements can be in a JSON file (array of {original, type, replacement}) or omitted for no replacements.
        input_path = Path(sys.argv[1])
        output_path = Path(sys.argv[2])
        replacements = []
        if len(sys.argv) >= 4:
            import json
            with open(sys.argv[3], encoding="utf-8") as f:
                replacements = json.load(f)
        if not input_path.is_file():
            print(f"File not found: {input_path}")
            sys.exit(1)
        file_bytes = input_path.read_bytes()
        out_bytes, mimetype = sanitize_document(file_bytes, input_path.name, replacements)
        output_path.write_bytes(out_bytes)
        print(f"Sanitized document saved to: {output_path} ({mimetype})")
    else:
        output = convert_to_text_file("input.pdf", ".")
        print(f"Saved to: {output}")